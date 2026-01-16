"""
Script para generar el Manual de Usuario de Gestión de Stocks
Formato: Microsoft Word (.docx)
Autor: Jobers
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from datetime import datetime
import os

# Colores corporativos
COLOR_PRIMARIO = RGBColor(0xFF, 0x43, 0x38)  # Rojo corporativo
COLOR_SECUNDARIO = RGBColor(0xD3, 0x2F, 0x2F)  # Rojo oscuro
COLOR_TEXTO = RGBColor(0x33, 0x33, 0x33)  # Gris oscuro
COLOR_FONDO_NOTA = RGBColor(0xE3, 0xF2, 0xFD)  # Azul claro para notas

def crear_estilos(doc):
    """Crear estilos corporativos para el documento"""
    styles = doc.styles

    # Estilo para Título 1
    style_h1 = styles['Heading 1']
    style_h1.font.name = 'Calibri'
    style_h1.font.size = Pt(24)
    style_h1.font.bold = True
    style_h1.font.color.rgb = COLOR_PRIMARIO
    style_h1.paragraph_format.space_before = Pt(24)
    style_h1.paragraph_format.space_after = Pt(12)

    # Estilo para Título 2
    style_h2 = styles['Heading 2']
    style_h2.font.name = 'Calibri'
    style_h2.font.size = Pt(18)
    style_h2.font.bold = True
    style_h2.font.color.rgb = COLOR_SECUNDARIO
    style_h2.paragraph_format.space_before = Pt(18)
    style_h2.paragraph_format.space_after = Pt(8)

    # Estilo para Título 3
    style_h3 = styles['Heading 3']
    style_h3.font.name = 'Calibri'
    style_h3.font.size = Pt(14)
    style_h3.font.bold = True
    style_h3.font.color.rgb = COLOR_TEXTO
    style_h3.paragraph_format.space_before = Pt(12)
    style_h3.paragraph_format.space_after = Pt(6)

    # Estilo Normal
    style_normal = styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXTO
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style_normal.paragraph_format.space_after = Pt(8)

def agregar_hipervinculo_interno(paragraph, text, bookmark_name):
    """Agregar un hipervínculo interno a un marcador"""
    # Crear el elemento de hipervínculo
    part = paragraph.part
    r_id = part.relate_to(
        'http://internal',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Estilo de hipervínculo
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FF4338')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def agregar_marcador(paragraph, bookmark_name):
    """Agregar un marcador a un párrafo"""
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    bookmark_start.set(qn('w:name'), bookmark_name)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(hash(bookmark_name) % 10000))

    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)

def agregar_nota(doc, tipo, texto):
    """Agregar una nota destacada (Consejo, Importante, Recomendación)"""
    iconos = {
        'consejo': '💡',
        'importante': '⚠️',
        'recomendacion': '✅',
        'nota': '📝'
    }

    colores = {
        'consejo': RGBColor(0xE8, 0xF5, 0xE9),  # Verde claro
        'importante': RGBColor(0xFF, 0xEB, 0xEE),  # Rojo claro
        'recomendacion': RGBColor(0xE3, 0xF2, 0xFD),  # Azul claro
        'nota': RGBColor(0xFF, 0xF8, 0xE1)  # Amarillo claro
    }

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    # Aplicar color de fondo
    shading = OxmlElement('w:shd')
    color_hex = '{:02X}{:02X}{:02X}'.format(*colores.get(tipo, colores['nota']))
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

    # Contenido de la nota
    icono = iconos.get(tipo, '📝')
    titulo_tipo = tipo.upper()

    p = cell.paragraphs[0]
    run = p.add_run(f"{icono} {titulo_tipo}: ")
    run.bold = True
    run.font.size = Pt(11)

    run2 = p.add_run(texto)
    run2.font.size = Pt(11)

    doc.add_paragraph()

def agregar_imagen_ficticia(doc, numero, titulo, descripcion_elementos):
    """Agregar descripción de una captura de pantalla ficticia"""
    # Marco para la imagen ficticia
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    # Borde
    set_cell_border(cell)

    # Contenido placeholder
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[Figura {numero}: {titulo}]\n\n")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    run.italic = True

    run2 = p.add_run("Captura de pantalla representativa\n")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Pie de imagen
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Figura {numero}: {titulo}")
    run_pie.bold = True
    run_pie.font.size = Pt(10)

    # Descripción de elementos
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_desc = desc.add_run(f"Elementos visibles: {descripcion_elementos}")
    run_desc.font.size = Pt(9)
    run_desc.italic = True
    run_desc.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

def set_cell_border(cell, border_color="CCCCCC", border_size="4"):
    """Establecer bordes de una celda"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)

    tcPr.append(tcBorders)

def agregar_tabla_simple(doc, headers, rows):
    """Agregar una tabla con formato corporativo"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Fondo rojo corporativo
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'FF4338')
        cell._tc.get_or_add_tcPr().append(shading)

    # Datos
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.rows[i+1].cells[j].text = str(value)

    doc.add_paragraph()

def crear_portada(doc):
    """Crear la portada del manual"""
    # Espacio superior
    for _ in range(3):
        doc.add_paragraph()

    # Título principal
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("GESTIÓN DE STOCKS")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARIO

    # Subtítulo
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Sistema de Consulta de Inventario en Tiempo Real")
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_TEXTO

    # Espacio
    for _ in range(2):
        doc.add_paragraph()

    # Línea decorativa
    linea = doc.add_paragraph()
    linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = linea.add_run("═" * 40)
    run.font.color.rgb = COLOR_PRIMARIO

    # Manual de Usuario
    manual = doc.add_paragraph()
    manual.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = manual.add_run("MANUAL DE USUARIO")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECUNDARIO

    # Línea decorativa
    linea2 = doc.add_paragraph()
    linea2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = linea2.add_run("═" * 40)
    run.font.color.rgb = COLOR_PRIMARIO

    # Espacio
    for _ in range(4):
        doc.add_paragraph()

    # Versión y fecha
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run(f"Versión 1.0")
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_TEXTO

    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fecha.add_run(datetime.now().strftime("%B %Y").title())
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_TEXTO

    # Salto de página
    doc.add_page_break()

def crear_indice(doc):
    """Crear índice del documento"""
    titulo = doc.add_heading('ÍNDICE', level=1)
    agregar_marcador(titulo, 'indice')

    # Secciones del índice
    secciones = [
        ("1. Introducción", "sec_introduccion"),
        ("2. Requisitos de Acceso", "sec_requisitos"),
        ("3. Visión General de la Aplicación", "sec_vision"),
        ("4. Guía Paso a Paso de Uso", "sec_guia"),
        ("5. Flujo del Proceso", "sec_flujo"),
        ("6. Casos de Uso Habituales", "sec_casos"),
        ("7. Errores Comunes y Resolución", "sec_errores"),
        ("8. Buenas Prácticas", "sec_buenas"),
        ("9. Preguntas Frecuentes (FAQ)", "sec_faq"),
        ("10. Conclusión", "sec_conclusion"),
    ]

    for texto, marcador in secciones:
        p = doc.add_paragraph()
        agregar_hipervinculo_interno(p, texto, marcador)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

def crear_seccion_introduccion(doc):
    """Crear sección de Introducción"""
    titulo = doc.add_heading('1. Introducción', level=1)
    agregar_marcador(titulo, 'sec_introduccion')

    # Qué es la aplicación
    doc.add_heading('1.1 ¿Qué es Gestión de Stocks?', level=2)
    doc.add_paragraph(
        "Gestión de Stocks es una aplicación web diseñada para la consulta y gestión de inventario "
        "de productos cerámicos y azulejos en tiempo real. Permite a los usuarios visualizar la "
        "disponibilidad de productos, consultar características técnicas, y realizar solicitudes "
        "de material de forma eficiente."
    )

    doc.add_paragraph(
        "La aplicación está desarrollada como una Single Page Application (SPA) con un backend "
        "REST API, lo que garantiza una experiencia de usuario fluida y rápida."
    )

    # Qué problema resuelve
    doc.add_heading('1.2 ¿Qué Problema Logístico Resuelve?', level=2)
    doc.add_paragraph(
        "En el sector cerámico, la gestión de inventario presenta desafíos únicos debido a:"
    )

    problemas = [
        "Variabilidad de tonos y calibres dentro de un mismo producto",
        "Necesidad de consultar existencias en tiempo real",
        "Coordinación entre múltiples almacenes y empresas",
        "Gestión de solicitudes y propuestas de material",
        "Trazabilidad de consultas y pedidos"
    ]

    for problema in problemas:
        p = doc.add_paragraph(problema, style='List Bullet')

    doc.add_paragraph(
        "Gestión de Stocks resuelve estos problemas proporcionando una plataforma centralizada "
        "donde los usuarios pueden consultar el stock disponible con detalle de tono, calibre y "
        "ubicación, realizar solicitudes y hacer seguimiento de las mismas."
    )

    # A quién va dirigida
    doc.add_heading('1.3 ¿A Quién Va Dirigida?', level=2)

    agregar_tabla_simple(doc,
        ['Tipo de Usuario', 'Descripción', 'Funciones Principales'],
        [
            ['Usuario Estándar', 'Comerciales, delegados de ventas', 'Consultar stock, crear solicitudes'],
            ['Administrador', 'Responsables de almacén', 'Gestionar usuarios, configurar sistema'],
            ['Superusuario', 'Administrador de sistemas', 'Acceso total, configuración avanzada']
        ]
    )

    agregar_nota(doc, 'consejo',
        "Si no está seguro de su tipo de usuario, consulte con el administrador del sistema.")

    # Enlace volver al índice
    p = doc.add_paragraph()
    agregar_hipervinculo_interno(p, "← Volver al índice", "indice")

    doc.add_page_break()

def crear_seccion_requisitos(doc):
    """Crear sección de Requisitos de Acceso"""
    titulo = doc.add_heading('2. Requisitos de Acceso', level=1)
    agregar_marcador(titulo, 'sec_requisitos')

    # Navegador recomendado
    doc.add_heading('2.1 Navegador Recomendado', level=2)
    doc.add_paragraph(
        "Para una experiencia óptima, se recomienda utilizar uno de los siguientes navegadores "
        "en su versión más reciente:"
    )

    agregar_tabla_simple(doc,
        ['Navegador', 'Versión Mínima', 'Recomendado'],
        [
            ['Google Chrome', '90+', '✅ Sí'],
            ['Mozilla Firefox', '88+', '✅ Sí'],
            ['Microsoft Edge', '90+', '✅ Sí'],
            ['Safari', '14+', '⚠️ Compatible'],
            ['Internet Explorer', 'No soportado', '❌ No']
        ]
    )

    agregar_nota(doc, 'importante',
        "Internet Explorer no es compatible. Si utiliza este navegador, actualice a Microsoft Edge.")

    # Dispositivos
    doc.add_heading('2.2 Uso en Diferentes Dispositivos', level=2)
    doc.add_paragraph(
        "La aplicación es completamente responsive y puede utilizarse en:"
    )

    dispositivos = [
        ("Ordenador (PC/Mac)", "Experiencia completa con todas las funcionalidades"),
        ("Tablet", "Interfaz adaptada con vista de tarjetas optimizada"),
        ("Móvil", "Diseño móvil-first con navegación simplificada")
    ]

    for dispositivo, descripcion in dispositivos:
        p = doc.add_paragraph()
        run = p.add_run(f"• {dispositivo}: ")
        run.bold = True
        p.add_run(descripcion)

    # Credenciales
    doc.add_heading('2.3 Credenciales de Acceso', level=2)
    doc.add_paragraph(
        "Para acceder al sistema necesita:"
    )

    credenciales = [
        "Nombre de usuario (proporcionado por el administrador)",
        "Contraseña (se le solicitará cambiarla en el primer acceso)",
        "Código de empresa (incluido en la URL de acceso)"
    ]

    for cred in credenciales:
        doc.add_paragraph(cred, style='List Bullet')

    agregar_nota(doc, 'importante',
        "La URL de acceso incluye el parámetro de empresa: https://servidor/login?empresa=X. "
        "Este parámetro es obligatorio para el correcto funcionamiento del sistema.")

    agregar_imagen_ficticia(doc, 1, "Pantalla de Login",
        "Logo de la empresa, campos de usuario y contraseña, botón de acceso, selector de idioma")

    # Enlace volver al índice
    p = doc.add_paragraph()
    agregar_hipervinculo_interno(p, "← Volver al índice", "indice")

    doc.add_page_break()

def crear_seccion_vision_general(doc):
    """Crear sección de Visión General"""
    titulo = doc.add_heading('3. Visión General de la Aplicación', level=1)
    agregar_marcador(titulo, 'sec_vision')

    # Pantalla principal
    doc.add_heading('3.1 Pantalla Principal', level=2)
    doc.add_paragraph(
        "Al acceder al sistema, se muestra la pantalla principal de consulta de stocks. "
        "Esta pantalla está dividida en las siguientes áreas:"
    )

    agregar_imagen_ficticia(doc, 2, "Pantalla Principal de Stocks",
        "Header con logo y menú de usuario, panel de filtros, tabla de productos, "
        "botón flotante de carrito, selector de idioma")

    areas = [
        ("Header", "Contiene el logo de la empresa, título de la aplicación y menú de usuario"),
        ("Panel de Filtros", "Permite filtrar productos por descripción, formato, serie, calidad, etc."),
        ("Tabla de Productos", "Muestra el listado de productos con sus características"),
        ("Botón de Carrito", "Acceso rápido al carrito de solicitudes (esquina inferior derecha)")
    ]

    for area, descripcion in areas:
        p = doc.add_paragraph()
        run = p.add_run(f"• {area}: ")
        run.bold = True
        p.add_run(descripcion)

    # Menús
    doc.add_heading('3.2 Menús y Navegación', level=2)
    doc.add_paragraph(
        "La aplicación ofrece diferentes opciones de menú según el rol del usuario:"
    )

    doc.add_heading('3.2.1 Menú de Usuario Estándar', level=3)
    menu_usuario = [
        "Stocks - Consulta de inventario",
        "Mis Propuestas - Historial de solicitudes realizadas",
        "Apariencia - Configuración de modo oscuro"
    ]
    for item in menu_usuario:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2.2 Menú de Administrador', level=3)
    menu_admin = [
        "Dashboard - Panel de estadísticas",
        "Todas las Propuestas - Gestión de solicitudes",
        "Todas las Consultas - Gestión de consultas",
        "Usuarios - Administración de usuarios",
        "Configuración Email - Configuración SMTP",
        "Parámetros - Parámetros del sistema",
        "Logo de Empresa - Personalización visual"
    ]
    for item in menu_admin:
        doc.add_paragraph(item, style='List Bullet')

    # Conceptos clave
    doc.add_heading('3.3 Conceptos Clave', level=2)

    conceptos = [
        ("Stock", "Cantidad disponible de un producto en el almacén"),
        ("Tono", "Variación de color dentro de un mismo producto"),
        ("Calibre", "Variación dimensional del producto"),
        ("Pallet", "Unidad de embalaje mayor (múltiples cajas)"),
        ("Caja", "Unidad de embalaje estándar"),
        ("Propuesta", "Solicitud de material enviada por un usuario"),
        ("Consulta", "Petición de información sobre un producto específico")
    ]

    agregar_tabla_simple(doc,
        ['Concepto', 'Definición'],
        [[c[0], c[1]] for c in conceptos]
    )

    # Enlace volver al índice
    p = doc.add_paragraph()
    agregar_hipervinculo_interno(p, "← Volver al índice", "indice")

    doc.add_page_break()

def crear_seccion_guia_paso_a_paso(doc):
    """Crear sección de Guía Paso a Paso"""
    titulo = doc.add_heading('4. Guía Paso a Paso de Uso', level=1)
    agregar_marcador(titulo, 'sec_guia')

    # Iniciar sesión
    doc.add_heading('4.1 Iniciar Sesión', level=2)
    pasos_login = [
        "Acceda a la URL proporcionada (incluye el parámetro de empresa)",
        "Introduzca su nombre de usuario en el campo correspondiente",
        "Introduzca su contraseña",
        "Haga clic en el botón 'Iniciar Sesión'",
        "Si es su primer acceso, se le solicitará cambiar la contraseña"
    ]
    for i, paso in enumerate(pasos_login, 1):
        doc.add_paragraph(f"{i}. {paso}")

    agregar_nota(doc, 'consejo',
        "Puede seleccionar el idioma de la interfaz antes de iniciar sesión usando el selector en la parte superior.")

    # Consultar stock
    doc.add_heading('4.2 Consultar Stock', level=2)
    doc.add_paragraph(
        "Para consultar la disponibilidad de productos:"
    )

    pasos_stock = [
        "En la pantalla principal, utilice los filtros disponibles (descripción, formato, serie, etc.)",
        "Haga clic en 'Buscar' para aplicar los filtros",
        "Los resultados se mostrarán en la tabla central",
        "Haga clic en cualquier fila para ver el detalle del producto",
        "En el detalle puede ver imágenes, ficha técnica y realizar consultas"
    ]
    for i, paso in enumerate(pasos_stock, 1):
        doc.add_paragraph(f"{i}. {paso}")

    agregar_imagen_ficticia(doc, 3, "Filtros de Búsqueda",
        "Campos de filtro: Descripción, Formato, Serie, Calidad, Color, botones Buscar y Limpiar")

    # Agregar al carrito
    doc.add_heading('4.3 Agregar Productos al Carrito', level=2)
    pasos_carrito = [
        "Localice el producto deseado en la tabla de resultados",
        "Haga clic en el botón 'Agregar' (icono de carrito) en la fila del producto",
        "En el modal que aparece, verifique los datos del producto",
        "Introduzca la cantidad deseada (puede usar los botones +/- o escribir directamente)",
        "Seleccione si desea la cantidad en cajas o pallets",
        "Haga clic en 'Agregar al Carrito'"
    ]
    for i, paso in enumerate(pasos_carrito, 1):
        doc.add_paragraph(f"{i}. {paso}")

    agregar_nota(doc, 'importante',
        "El sistema detecta duplicados automáticamente. Si agrega el mismo producto con idénticas "
        "características (código, calidad, tono, calibre, pallet, caja), se le notificará.")

    # Enviar propuesta
    doc.add_heading('4.4 Enviar una Propuesta', level=2)
    pasos_propuesta = [
        "Haga clic en el botón flotante del carrito (esquina inferior derecha)",
        "Revise los productos añadidos al carrito",
        "Puede eliminar productos individuales si lo desea",
        "Haga clic en 'Enviar Propuesta'",
        "Complete el formulario de envío (nombre, email, comentarios)",
        "Opcionalmente, marque la casilla para recibir una copia del email",
        "Haga clic en 'Enviar'"
    ]
    for i, paso in enumerate(pasos_propuesta, 1):
        doc.add_paragraph(f"{i}. {paso}")

    agregar_imagen_ficticia(doc, 4, "Modal del Carrito",
        "Lista de productos añadidos, cantidad solicitada, botones de eliminar, "
        "total de items, botón Enviar Propuesta")

    # Consultar producto
    doc.add_heading('4.5 Realizar una Consulta sobre un Producto', level=2)
    pasos_consulta = [
        "Haga clic en un producto para abrir el detalle",
        "En la sección 'Contacto', haga clic en 'Consultar'",
        "Complete el formulario con sus datos y la consulta",
        "Haga clic en 'Enviar Consulta'",
        "Alternativamente, puede usar el botón de WhatsApp para contacto directo"
    ]
    for i, paso in enumerate(pasos_consulta, 1):
        doc.add_paragraph(f"{i}. {paso}")

    # Enlace volver al índice
    p = doc.add_paragraph()
    agregar_hipervinculo_interno(p, "← Volver al índice", "indice")

    doc.add_page_break()

def crear_seccion_flujo(doc):
    """Crear sección de Flujo del Proceso"""
    titulo = doc.add_heading('5. Flujo del Proceso', level=1)
    agregar_marcador(titulo, 'sec_flujo')

    doc.add_heading('5.1 Diagrama del Proceso de Solicitud', level=2)
    doc.add_paragraph(
        "El siguiente diagrama muestra el flujo completo desde la consulta hasta la gestión de la propuesta:"
    )

    # Diagrama como tabla
    flujo = [
        "┌─────────────────┐",
        "│   INICIO        │",
        "│   Login         │",
        "└────────┬────────┘",
        "         │",
        "         ▼",
        "┌─────────────────┐",
        "│ Consultar Stock │",
        "│ (Filtrar)       │",
        "└────────┬────────┘",
        "         │",
        "         ▼",
        "┌─────────────────┐",
        "│ Seleccionar     │",
        "│ Productos       │",
        "└────────┬────────┘",
        "         │",
        "         ▼",
        "┌─────────────────┐",
        "│ Agregar al      │",
        "│ Carrito         │",
        "└────────┬────────┘",
        "         │",
        "         ▼",
        "┌─────────────────┐",
        "│ Revisar         │",
        "│ Carrito         │",
        "└────────┬────────┘",
        "         │",
        "         ▼",
        "┌─────────────────┐",
        "│ Enviar          │",
        "│ Propuesta       │",
        "└────────┬────────┘",
        "         │",
        "         ▼",
        "┌─────────────────┐",
        "│ Propuesta       │",
        "│ Enviada ✓       │",
        "└─────────────────┘",
    ]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for linea in flujo:
        run = p.add_run(linea + "\n")
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    doc.add_heading('5.2 Estados de una Propuesta', level=2)

    agregar_tabla_simple(doc,
        ['Estado', 'Descripción', 'Acción Requerida'],
        [
            ['Enviada', 'Propuesta recibida, pendiente de procesar', 'Esperar procesamiento'],
            ['Procesada', 'Propuesta gestionada por administración', 'Ninguna'],
            ['Cancelada', 'Propuesta cancelada', 'Contactar si es necesario']
        ]
    )

    doc.add_heading('5.3 Diagrama de Navegación', level=2)
    doc.add_paragraph(
        "Esquema de navegación entre las principales pantallas de la aplicación:"
    )

    navegacion = """
    ┌──────────────────────────────────────────────────────────────┐
    │                         APLICACIÓN                          │
    ├──────────────────────────────────────────────────────────────┤
    │                                                              │
    │   ┌─────────┐      ┌─────────────┐      ┌──────────────┐   │
    │   │  Login  │ ───► │   Stocks    │ ───► │   Detalle    │   │
    │   └─────────┘      │  (Principal)│      │   Producto   │   │
    │                    └──────┬──────┘      └──────────────┘   │
    │                           │                                 │
    │              ┌────────────┼────────────┐                   │
    │              ▼            ▼            ▼                   │
    │        ┌──────────┐ ┌──────────┐ ┌──────────────┐         │
    │        │ Carrito  │ │   Mis    │ │ Administrac. │         │
    │        │          │ │Propuestas│ │   (Admin)    │         │
    │        └──────────┘ └──────────┘ └──────────────┘         │
    │                                                              │
    └──────────────────────────────────────────────────────────────┘
    """

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(navegacion)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    # Enlace volver al índice
    p = doc.add_paragraph()
    agregar_hipervinculo_interno(p, "← Volver al índice", "indice")

    doc.add_page_break()

def crear_seccion_casos_uso(doc):
    """Crear sección de Casos de Uso"""
    titulo = doc.add_heading('6. Casos de Uso Habituales', level=1)
    agregar_marcador(titulo, 'sec_casos')

    # Caso 1
    doc.add_heading('6.1 Caso: Comercial Consultando Disponibilidad', level=2)
    doc.add_paragraph(
        "Escenario: Un comercial necesita verificar si hay stock disponible de un azulejo "
        "específico para un cliente."
    )

    pasos = [
        "El comercial accede al sistema con sus credenciales",
        "En el filtro de descripción, escribe el nombre del producto",
        "Hace clic en 'Buscar'",
        "Revisa las existencias en la columna correspondiente",
        "Si hay stock, anota el tono y calibre disponibles",
        "Opcionalmente, descarga la ficha técnica para enviar al cliente"
    ]

    doc.add_paragraph("Pasos:", style='Heading 3')
    for i, paso in enumerate(pasos, 1):
        doc.add_paragraph(f"{i}. {paso}")

    agregar_nota(doc, 'consejo',
        "Utilice el filtro de 'Existencias mínimas' para mostrar solo productos con stock disponible.")

    # Caso 2
    doc.add_heading('6.2 Caso: Crear Solicitud de Material', level=2)
    doc.add_paragraph(
        "Escenario: El comercial necesita solicitar material para un proyecto."
    )

    pasos2 = [
        "Busca los productos necesarios usando los filtros",
        "Para cada producto, hace clic en 'Agregar' y especifica la cantidad",
        "Una vez añadidos todos los productos, abre el carrito",
        "Revisa que las cantidades sean correctas",
        "Hace clic en 'Enviar Propuesta'",
        "Completa el formulario indicando el proyecto/cliente",
        "Envía la propuesta"
    ]

    doc.add_paragraph("Pasos:", style='Heading 3')
    for i, paso in enumerate(pasos2, 1):
        doc.add_paragraph(f"{i}. {paso}")

    # Caso 3
    doc.add_heading('6.3 Caso: Administrador Gestionando Propuestas', level=2)
    doc.add_paragraph(
        "Escenario: El administrador necesita revisar y procesar las propuestas pendientes."
    )

    pasos3 = [
        "Accede al sistema con credenciales de administrador",
        "En el menú, selecciona 'Todas las Propuestas'",
        "Filtra por estado 'Enviada' para ver las pendientes",
        "Hace clic en una propuesta para ver el detalle",
        "Revisa los productos solicitados",
        "Cambia el estado a 'Procesada' cuando se gestione"
    ]

    doc.add_paragraph("Pasos:", style='Heading 3')
    for i, paso in enumerate(pasos3, 1):
        doc.add_paragraph(f"{i}. {paso}")

    # Caso 4
    doc.add_heading('6.4 Caso: Consulta Técnica sobre un Producto', level=2)
    doc.add_paragraph(
        "Escenario: Un cliente necesita información técnica específica de un producto."
    )

    pasos4 = [
        "El comercial busca el producto en el sistema",
        "Hace clic en el producto para abrir el detalle",
        "Si existe ficha técnica, hace clic en 'Descargar PDF'",
        "Si necesita información adicional, usa el botón 'Consultar'",
        "Completa el formulario describiendo la consulta",
        "Envía la consulta, que llegará al departamento correspondiente"
    ]

    doc.add_paragraph("Pasos:", style='Heading 3')
    for i, paso in enumerate(pasos4, 1):
        doc.add_paragraph(f"{i}. {paso}")

    # Enlace volver al índice
    p = doc.add_paragraph()
    agregar_hipervinculo_interno(p, "← Volver al índice", "indice")

    doc.add_page_break()

def crear_seccion_errores(doc):
    """Crear sección de Errores Comunes"""
    titulo = doc.add_heading('7. Errores Comunes y Resolución', level=1)
    agregar_marcador(titulo, 'sec_errores')

    errores = [
        {
            "error": "Error: 'Falta el parámetro empresa'",
            "causa": "La URL de acceso no incluye el parámetro de empresa obligatorio",
            "solucion": "Asegúrese de usar la URL completa proporcionada: https://servidor/login?empresa=X"
        },
        {
            "error": "No se puede iniciar sesión",
            "causa": "Credenciales incorrectas o usuario inactivo",
            "solucion": "Verifique usuario/contraseña. Si persiste, contacte al administrador"
        },
        {
            "error": "La página no carga correctamente",
            "causa": "Navegador no compatible o caché corrupta",
            "solucion": "Actualice el navegador o limpie la caché (Ctrl+Shift+Delete)"
        },
        {
            "error": "No aparecen productos en la búsqueda",
            "causa": "Filtros demasiado restrictivos o sin stock",
            "solucion": "Limpie los filtros con el botón 'Limpiar' y busque de nuevo"
        },
        {
            "error": "Error al enviar propuesta",
            "causa": "Carrito vacío o problema de conexión",
            "solucion": "Verifique que hay productos en el carrito. Reintente o recargue la página"
        },
        {
            "error": "No se ve el botón del carrito",
            "causa": "Las propuestas están deshabilitadas para su empresa",
            "solucion": "Contacte al administrador para habilitar la funcionalidad"
        },
        {
            "error": "Las imágenes no cargan",
            "causa": "Conexión lenta o producto sin imágenes",
            "solucion": "Espere a que carguen o verifique su conexión a internet"
        },
        {
            "error": "El modo oscuro no se aplica",
            "causa": "Problema con localStorage del navegador",
            "solucion": "Limpie los datos del sitio en la configuración del navegador"
        }
    ]

    for i, error in enumerate(errores, 1):
        doc.add_heading(f'7.{i} {error["error"]}', level=2)

        p = doc.add_paragraph()
        run = p.add_run("Causa: ")
        run.bold = True
        p.add_run(error["causa"])

        p2 = doc.add_paragraph()
        run2 = p2.add_run("Solución: ")
        run2.bold = True
        p2.add_run(error["solucion"])

    agregar_nota(doc, 'importante',
        "Si experimenta un error no listado aquí, tome una captura de pantalla del mensaje "
        "y contacte al soporte técnico.")

    # Enlace volver al índice
    p = doc.add_paragraph()
    agregar_hipervinculo_interno(p, "← Volver al índice", "indice")

    doc.add_page_break()

def crear_seccion_buenas_practicas(doc):
    """Crear sección de Buenas Prácticas"""
    titulo = doc.add_heading('8. Buenas Prácticas', level=1)
    agregar_marcador(titulo, 'sec_buenas')

    doc.add_heading('8.1 Seguridad', level=2)
    practicas_seguridad = [
        "Cambie su contraseña periódicamente (cada 90 días recomendado)",
        "No comparta sus credenciales con otros usuarios",
        "Cierre sesión al terminar de usar la aplicación",
        "No acceda desde ordenadores públicos o no seguros",
        "Reporte inmediatamente cualquier acceso sospechoso"
    ]
    for practica in practicas_seguridad:
        doc.add_paragraph(practica, style='List Bullet')

    doc.add_heading('8.2 Uso Eficiente', level=2)
    practicas_uso = [
        "Utilice filtros específicos para reducir el tiempo de búsqueda",
        "Guarde las URLs de acceso frecuente en favoritos",
        "Revise el carrito antes de enviar para evitar errores",
        "Añada comentarios descriptivos en las propuestas",
        "Consulte el historial de propuestas antes de crear nuevas"
    ]
    for practica in practicas_uso:
        doc.add_paragraph(practica, style='List Bullet')

    doc.add_heading('8.3 Organización', level=2)
    practicas_org = [
        "Agrupe los productos del mismo proyecto en una sola propuesta",
        "Indique claramente el cliente/proyecto en los comentarios",
        "Revise las propuestas pendientes regularmente",
        "Mantenga actualizada su información de contacto"
    ]
    for practica in practicas_org:
        doc.add_paragraph(practica, style='List Bullet')

    agregar_nota(doc, 'recomendacion',
        "Configure el modo oscuro si trabaja en ambientes con poca luz. "
        "Reduce la fatiga visual y mejora la experiencia de uso.")

    # Enlace volver al índice
    p = doc.add_paragraph()
    agregar_hipervinculo_interno(p, "← Volver al índice", "indice")

    doc.add_page_break()

def crear_seccion_faq(doc):
    """Crear sección de Preguntas Frecuentes"""
    titulo = doc.add_heading('9. Preguntas Frecuentes (FAQ)', level=1)
    agregar_marcador(titulo, 'sec_faq')

    faqs = [
        {
            "pregunta": "¿Puedo acceder desde mi móvil?",
            "respuesta": "Sí, la aplicación es completamente responsive y funciona en cualquier dispositivo con navegador web moderno."
        },
        {
            "pregunta": "¿Cómo cambio mi contraseña?",
            "respuesta": "Actualmente, el cambio de contraseña se realiza a través del administrador del sistema. Contacte con él para solicitar el cambio."
        },
        {
            "pregunta": "¿Qué significan los diferentes tonos y calibres?",
            "respuesta": "El tono indica variaciones de color dentro del mismo producto. El calibre indica variaciones de tamaño. Ambos son importantes para garantizar uniformidad en una instalación."
        },
        {
            "pregunta": "¿Puedo cancelar una propuesta enviada?",
            "respuesta": "Una vez enviada, la propuesta no puede cancelarse por el usuario. Contacte al administrador si necesita cancelarla."
        },
        {
            "pregunta": "¿Por qué no veo ciertos productos?",
            "respuesta": "Los productos mostrados dependen de la empresa configurada. Si no encuentra un producto, verifique que está accediendo con la empresa correcta o contacte al administrador."
        },
        {
            "pregunta": "¿Cómo descargo la ficha técnica de un producto?",
            "respuesta": "Abra el detalle del producto haciendo clic en él. Si existe ficha técnica, verá un botón 'Descargar PDF' en la sección correspondiente."
        },
        {
            "pregunta": "¿Puedo cambiar el idioma de la aplicación?",
            "respuesta": "Sí, use el selector de idioma en la parte superior de la pantalla. Están disponibles: Español, Inglés y Francés."
        },
        {
            "pregunta": "¿Qué hago si olvidé mi contraseña?",
            "respuesta": "Contacte al administrador del sistema para que le genere una nueva contraseña temporal."
        },
        {
            "pregunta": "¿Las propuestas tienen fecha de validez?",
            "respuesta": "La validez depende de las políticas de su empresa. Consulte con su responsable comercial."
        },
        {
            "pregunta": "¿Puedo ver propuestas de otros usuarios?",
            "respuesta": "Solo los administradores pueden ver todas las propuestas. Los usuarios estándar solo ven sus propias propuestas."
        }
    ]

    for i, faq in enumerate(faqs, 1):
        # Pregunta
        p = doc.add_paragraph()
        run = p.add_run(f"P{i}: {faq['pregunta']}")
        run.bold = True
        run.font.color.rgb = COLOR_PRIMARIO

        # Respuesta
        p2 = doc.add_paragraph()
        run2 = p2.add_run("R: ")
        run2.bold = True
        p2.add_run(faq['respuesta'])

        doc.add_paragraph()

    # Enlace volver al índice
    p = doc.add_paragraph()
    agregar_hipervinculo_interno(p, "← Volver al índice", "indice")

    doc.add_page_break()

def crear_seccion_conclusion(doc):
    """Crear sección de Conclusión"""
    titulo = doc.add_heading('10. Conclusión', level=1)
    agregar_marcador(titulo, 'sec_conclusion')

    doc.add_paragraph(
        "Gestión de Stocks es una herramienta diseñada para optimizar el proceso de consulta "
        "y gestión de inventario en el sector cerámico. Su interfaz intuitiva y sus funcionalidades "
        "avanzadas permiten a los usuarios trabajar de forma eficiente y organizada."
    )

    doc.add_heading('Resumen de Funcionalidades Principales', level=2)
    funcionalidades = [
        "Consulta de stock en tiempo real con filtros avanzados",
        "Sistema de solicitudes (propuestas) con seguimiento",
        "Visualización de imágenes y fichas técnicas",
        "Consultas directas sobre productos",
        "Panel de administración completo",
        "Soporte multi-idioma (ES, EN, FR)",
        "Modo oscuro para mayor confort visual",
        "Diseño responsive para cualquier dispositivo"
    ]
    for func in funcionalidades:
        doc.add_paragraph(func, style='List Bullet')

    doc.add_heading('Soporte Técnico', level=2)
    doc.add_paragraph(
        "Para cualquier consulta o incidencia técnica, contacte con:"
    )

    soporte = [
        ("Email", "soporte@empresa.com"),
        ("Teléfono", "+34 XXX XXX XXX"),
        ("Horario", "Lunes a Viernes, 9:00 - 18:00")
    ]

    for campo, valor in soporte:
        p = doc.add_paragraph()
        run = p.add_run(f"• {campo}: ")
        run.bold = True
        p.add_run(valor)

    agregar_nota(doc, 'nota',
        "Este manual se actualiza periódicamente. Consulte la versión más reciente "
        "en caso de duda sobre alguna funcionalidad.")

    # Agradecimiento final
    doc.add_paragraph()
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final.add_run("Gracias por utilizar Gestión de Stocks")
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_PRIMARIO
    run.bold = True

    # Enlace volver al índice
    doc.add_paragraph()
    p = doc.add_paragraph()
    agregar_hipervinculo_interno(p, "← Volver al índice", "indice")

def configurar_pie_pagina(doc):
    """Configurar pie de página con numeración"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Texto del pie
    run = p.add_run("Gestión de Stocks - Manual de Usuario | Página ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    # Número de página (campo)
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.text = 'PAGE'

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run2 = p.add_run()
    run2._r.append(fld_char_begin)
    run2._r.append(instr_text)
    run2._r.append(fld_char_end)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

def main():
    """Función principal para generar el manual"""
    print("Generando Manual de Usuario de Gestión de Stocks...")

    # Crear documento
    doc = Document()

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Crear estilos
    crear_estilos(doc)

    # Crear secciones del documento
    print("  - Creando portada...")
    crear_portada(doc)

    print("  - Creando índice...")
    crear_indice(doc)

    print("  - Creando sección de introducción...")
    crear_seccion_introduccion(doc)

    print("  - Creando sección de requisitos...")
    crear_seccion_requisitos(doc)

    print("  - Creando visión general...")
    crear_seccion_vision_general(doc)

    print("  - Creando guía paso a paso...")
    crear_seccion_guia_paso_a_paso(doc)

    print("  - Creando flujo del proceso...")
    crear_seccion_flujo(doc)

    print("  - Creando casos de uso...")
    crear_seccion_casos_uso(doc)

    print("  - Creando sección de errores...")
    crear_seccion_errores(doc)

    print("  - Creando buenas prácticas...")
    crear_seccion_buenas_practicas(doc)

    print("  - Creando FAQ...")
    crear_seccion_faq(doc)

    print("  - Creando conclusión...")
    crear_seccion_conclusion(doc)

    print("  - Configurando pie de página...")
    configurar_pie_pagina(doc)

    # Guardar documento
    output_path = os.path.join(os.path.dirname(__file__), 'Manual_Usuario_Gestion_Stocks.docx')
    doc.save(output_path)

    print(f"\n[OK] Manual generado exitosamente: {output_path}")
    print("\nEl documento incluye:")
    print("  - Portada profesional")
    print("  - Indice clicable")
    print("  - 10 secciones completas")
    print("  - Tablas con estilo corporativo")
    print("  - Notas y recomendaciones destacadas")
    print("  - Diagramas de flujo")
    print("  - FAQ con 10 preguntas")
    print("  - Pie de pagina con numeracion")
    print("\nPuede exportar a PDF desde Word: Archivo > Guardar como > PDF")

if __name__ == "__main__":
    main()
